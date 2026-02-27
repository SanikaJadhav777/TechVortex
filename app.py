import streamlit as st
from groq import Groq
from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# ------------------------------------------------
# PAGE CONFIG
# ------------------------------------------------
st.set_page_config(
    page_title="TechVortex",
    page_icon="💡",
    layout="wide"
)

# ------------------------------------------------
# GLOBAL CSS
# ------------------------------------------------
st.markdown("""
<style>
body {
    background: linear-gradient(135deg, #eef2f3, #e9e4f0);
}
[data-testid="stAppViewContainer"] {
    background: transparent;
}
header { visibility: hidden; }

.topbar {
    background: #3f51b5;
    color: white;
    padding: 14px 30px;
    font-size: 20px;
    font-weight: 600;
    display: flex;
    justify-content: space-between;
    align-items: center;
}

.card-header {
    background: linear-gradient(90deg, #1e88e5, #43a047);
    color: white;
    padding: 16px 20px;
    border-radius: 12px;
    font-size: 20px;
    font-weight: 600;
    display: flex;
    justify-content: space-between;
}

.counter {
    background: rgba(255,255,255,0.9);
    color: #333;
    border-radius: 20px;
    padding: 6px 14px;
    font-size: 13px;
}

div.stButton > button {
    background-color: #3f51b5 !important;
    color: white !important;
    border-radius: 20px !important;
    height: 42px !important;
    font-weight: 600 !important;
}
div.stButton > button:hover {
    background-color: #303f9f !important;
}
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------
# TOP BAR
# ------------------------------------------------
st.markdown("""
<div class="topbar">
<div>TechVortex</div>
<div>Logout</div>
</div>
""", unsafe_allow_html=True)

# ------------------------------------------------
# GROQ SETUP
# ------------------------------------------------
try:
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
except Exception:
    st.error("⚠ GROQ_API_KEY not configured.")
    st.stop()

# ------------------------------------------------
# SESSION STATE
# ------------------------------------------------
st.session_state.setdefault("text_key", 0)
st.session_state.setdefault("initial_story", None)
st.session_state.setdefault("chat_history", [])
st.session_state.setdefault("followup_input", "")
st.session_state.setdefault("draft", "")
st.session_state.setdefault("last_uploaded", None)

# ------------------------------------------------
# HELPERS
# ------------------------------------------------
def extract_text(file):
    text = ""
    if file.type == "application/pdf":
        reader = PdfReader(file)
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file.type == "text/plain":
        text = file.read().decode("utf-8")
    return text

def generate_initial_story(requirement, context):
    prompt = f"""
You are a Senior Agile Business Analyst.

Convert this requirement into:
- Atomic user stories
- Acceptance Criteria
- Edge Cases
- Assumptions

STRICT FORMAT.

Requirement:
{requirement}
"""
    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return resp.choices[0].message.content

def generate_followup(question):
    messages = [{"role": "assistant", "content": st.session_state.initial_story}]
    for h in st.session_state.chat_history:
        messages.append({"role": "assistant", "content": h})
    messages.append({"role": "user", "content": question})

    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=0.5
    )
    answer = resp.choices[0].message.content
    st.session_state.chat_history.append(answer)
    return answer

def clear_all():
    st.session_state.draft = ""
    st.session_state.initial_story = None
    st.session_state.chat_history = []
    st.session_state.followup_input = ""

# ------------------------------------------------
# DOCUMENT BUILDERS (DOWNLOAD)
# ------------------------------------------------
def build_word(content):
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_pdf(content):
    buf = BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    for line in content.split("\n"):
        if y < 40:
            pdf.showPage()
            y = height - 40
        pdf.drawString(x, y, line)
        y -= 14
    pdf.save()
    buf.seek(0)
    return buf

# ------------------------------------------------
# MAIN INPUT
# ------------------------------------------------
requirement = st.session_state.draft
words = len(requirement.split())
chars = len(requirement)

st.markdown(f"""
<div class="card-header">
<span>Provide Requirements</span>
<div>
<span class="counter">Words: {words}</span>
<span class="counter">Characters: {chars}</span>
</div>
</div>
""", unsafe_allow_html=True)

tab_text, tab_file = st.tabs(["Text", "File"])

with tab_file:
    uploaded_file = st.file_uploader("Upload file", type=["docx", "pdf", "txt"])
    if uploaded_file and st.session_state.last_uploaded != uploaded_file.name:
        st.session_state.draft = extract_text(uploaded_file)
        st.session_state.last_uploaded = uploaded_file.name

with tab_text:
    st.text_area(
        "Requirement",
        key="draft",
        height=220,
        label_visibility="collapsed"
    )

# ------------------------------------------------
# ACTION BUTTONS
# ------------------------------------------------
col1, col2, col3, col4, col5 = st.columns([1,1,1,3,1])

with col1:
    st.button("💾 Save Draft")

with col2:
    if st.button("🔄 Regenerate"):
        if st.session_state.draft.strip():
            st.session_state.initial_story = generate_initial_story(st.session_state.draft, "")
            st.session_state.chat_history = []

with col3:
    st.button("❌ Clear", on_click=clear_all)

with col5:
    if st.button("✨ Generate"):
        if st.session_state.draft.strip():
            st.session_state.initial_story = generate_initial_story(st.session_state.draft, "")

# ------------------------------------------------
# OUTPUT
# ------------------------------------------------
if st.session_state.initial_story:
    st.markdown("## 📝 Generated User Stories")
    st.markdown(st.session_state.initial_story)

    st.markdown("### ⬇️ Download Output")
    fmt = st.selectbox("Format", ["Word (.docx)", "PDF (.pdf)"], label_visibility="collapsed")

    if fmt == "Word (.docx)":
        st.download_button(
            "📄 Download Word",
            build_word(st.session_state.initial_story),
            file_name="Generated_User_Stories.docx"
        )
    else:
        st.download_button(
            "📕 Download PDF",
            build_pdf(st.session_state.initial_story),
            file_name="Generated_User_Stories.pdf"
        )

    st.markdown("## 💬 Follow-up Questions")
    st.session_state.followup_input = st.text_area(
        "Ask question",
        value=st.session_state.followup_input,
        height=100
    )

    if st.button("Ask AI"):
        if st.session_state.followup_input.strip():
            answer = generate_followup(st.session_state.followup_input)
            st.session_state.followup_input = ""
            st.markdown(answer)

if st.session_state.chat_history:
    st.markdown("## 🗂 Follow-up History")
    for i, h in enumerate(st.session_state.chat_history, 1):
        st.markdown(f"**Follow-up {i}:** {h}")
